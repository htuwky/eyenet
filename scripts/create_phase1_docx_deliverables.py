from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

DELIVERABLES = [
    (
        Path("docs/group_meeting_report_2026-05-28.md"),
        Path("docs/group_meeting_report_2026-05-28.docx"),
        "EyeNet phase-1 group meeting report",
    ),
    (
        Path("docs/paper_drafts/content_agnostic_eye_movement_screening_phase1_draft.md"),
        Path("docs/paper_drafts/content_agnostic_eye_movement_screening_phase1_draft.docx"),
        "EyeNet phase-1 paper draft",
    ),
]


def main() -> None:
    for source, output, footer_text in DELIVERABLES:
        doc = markdown_to_docx(source, footer_text=footer_text)
        output.parent.mkdir(parents=True, exist_ok=True)
        doc.save(output)
        print(output)


def markdown_to_docx(path: Path, footer_text: str) -> Document:
    doc = Document()
    configure_document(doc, footer_text)
    lines = path.read_text(encoding="utf-8").splitlines()
    index = 0
    in_code = False
    code_lines: list[str] = []

    while index < len(lines):
        line = lines[index]
        stripped = line.strip()

        if stripped.startswith("```"):
            if in_code:
                add_code_block(doc, code_lines)
                code_lines = []
                in_code = False
            else:
                in_code = True
            index += 1
            continue

        if in_code:
            code_lines.append(line)
            index += 1
            continue

        if not stripped:
            index += 1
            continue

        if is_table_start(lines, index):
            table_lines = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                table_lines.append(lines[index])
                index += 1
            add_markdown_table(doc, table_lines)
            continue

        if stripped.startswith("#"):
            add_heading(doc, stripped)
        elif re.match(r"^\d+\.\s+", stripped):
            add_list_item(doc, re.sub(r"^\d+\.\s+", "", stripped), numbered=True)
        elif stripped.startswith("- "):
            add_list_item(doc, stripped[2:], numbered=False)
        else:
            doc.add_paragraph(stripped)
        index += 1

    return doc


def configure_document(doc: Document, footer_text: str) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(5)

    title = doc.styles["Title"]
    title.font.name = "Microsoft YaHei"
    title._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    title.font.size = Pt(20)
    title.font.bold = True
    title.font.color.rgb = RGBColor(0x0B, 0x25, 0x45)

    for style_name, size, color in [
        ("Heading 1", 15, RGBColor(0x2E, 0x74, 0xB5)),
        ("Heading 2", 12.5, RGBColor(0x2E, 0x74, 0xB5)),
        ("Heading 3", 11.5, RGBColor(0x1F, 0x4D, 0x78)),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(9)
        style.paragraph_format.space_after = Pt(4)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run(f"{footer_text} | 2026-05-27")
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def add_heading(doc: Document, line: str) -> None:
    level = min(len(line) - len(line.lstrip("#")), 3)
    text = line.lstrip("#").strip()
    if level == 1 and not any(par.style.name == "Title" for par in doc.paragraphs):
        paragraph = doc.add_paragraph(style="Title")
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.add_run(text)
    else:
        doc.add_heading(text, level=level)


def add_list_item(doc: Document, text: str, numbered: bool) -> None:
    style = "List Number" if numbered else "List Bullet"
    paragraph = doc.add_paragraph(style=style)
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.add_run(text)


def add_code_block(doc: Document, lines: list[str]) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.18)
    paragraph.paragraph_format.space_before = Pt(3)
    paragraph.paragraph_format.space_after = Pt(6)
    for idx, line in enumerate(lines):
        run = paragraph.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(8.8)
        if idx < len(lines) - 1:
            run.add_break()


def is_table_start(lines: list[str], index: int) -> bool:
    return (
        index + 1 < len(lines)
        and lines[index].strip().startswith("|")
        and re.match(r"^\|\s*:?-{3,}:?\s*(\|\s*:?-{3,}:?\s*)+\|?$", lines[index + 1].strip()) is not None
    )


def add_markdown_table(doc: Document, table_lines: list[str]) -> None:
    rows = [parse_table_row(line) for line in table_lines]
    headers = rows[0]
    body = rows[2:]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = compute_widths(headers, body)
    set_table_geometry(table, widths)

    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = clean_inline_markdown(header)
        set_cell_shading(cell, "EAF2F8")
        format_cell(cell, bold=True)

    for row in body:
        cells = table.add_row().cells
        for idx, text in enumerate(row[: len(headers)]):
            cells[idx].text = clean_inline_markdown(text)
            format_cell(cells[idx], bold=False)

    doc.add_paragraph()


def parse_table_row(line: str) -> list[str]:
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def clean_inline_markdown(text: str) -> str:
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
    return text.replace("`", "")


def compute_widths(headers: list[str], body: list[list[str]]) -> list[int]:
    content = [headers, *body]
    scores = []
    for col_idx in range(len(headers)):
        score = max(len(row[col_idx]) if col_idx < len(row) else 1 for row in content)
        scores.append(max(score, 6))
    total = sum(scores)
    raw = [max(700, int(9360 * score / total)) for score in scores]
    raw[-1] += 9360 - sum(raw)
    return raw


def set_table_geometry(table, widths: list[int]) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_grid = tbl.tblGrid
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        tbl_grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[idx])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_margins(cell, top: int = 80, start: int = 100, bottom: int = 80, end: int = 100) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def format_cell(cell, *, bold: bool) -> None:
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_after = Pt(0)
        if paragraph.text and len(paragraph.text) < 16:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            run.bold = bold
            run.font.size = Pt(8.0)
            run.font.name = "Microsoft YaHei"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")


if __name__ == "__main__":
    main()
