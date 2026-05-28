from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

OUTPUT = Path("docs/paper_drafts/content_agnostic_eye_movement_screening_draft.docx")


DATASET_ROWS = [
    ["EMS", "Complete", "160", "225,159", "Primary SZ/HC downstream benchmark."],
    ["HBN", "Complete", "1,244 usable", "1,684,382", "Technically integrated; not currently the best transfer source."],
    ["GazeBase", "Complete", "322", "843,517", "Video tasks VD1/VD2; high-specificity single-split behavior."],
    ["OneStop", "Complete", "360", "2,042,834", "Reading corpus; technically integrated, weaker EMS transfer."],
    ["CRCNS eye-1", "Complete", "16", "67,172", "Natural movie-viewing source; current best public transfer route."],
    ["Saliency4ASD", "Deferred", "TBD", "TBD", "Pseudo-subject/session structure; ASD labels must not be merged with SZ/HC."],
]


EMS_BASELINE_ROWS = [
    ["From scratch", "0.784 +/- 0.069", "0.700 +/- 0.114", "0.800", "0.600", "Supervised encoder baseline."],
    ["EMS MEM frozen", "0.716 +/- 0.059", "0.619 +/- 0.081", "0.813", "0.425", "Pretrained encoder alone is insufficient."],
    ["EMS MEM fine-tune", "0.832 +/- 0.065", "0.763 +/- 0.052", "0.800", "0.725", "Best EMS-only multi-seed route."],
]


PUBLIC_SCREENING_ROWS = [
    ["EMS-only MEM", "0.891", "0.750", "0.938", "0.563", "Strong EMS-only seed42 result."],
    ["HBN+EMS MEM", "0.859", "0.719", "0.875", "0.563", "No clear gain over EMS-only."],
    ["GazeBase+EMS MEM", "0.863", "0.813", "0.750", "0.875", "High-specificity operating point."],
    ["OneStop+EMS MEM", "0.891", "0.719", "0.938", "0.500", "Good AUC, weak specificity."],
    ["CRCNS eye-1+EMS MEM", "0.910", "0.781", "0.750", "0.813", "Best public single-source AUC candidate."],
    ["GazeBase+CRCNS+EMS MEM", "0.898", "0.781", "0.938", "0.625", "Fusion did not clearly beat CRCNS-only."],
]


ALIGNED_MODEL_SELECTION_ROWS = [
    [
        "CRCNS+EMS MEM, dropout 0.3",
        "fine-tune",
        "5",
        "0.813 +/- 0.062",
        "0.725 +/- 0.060",
        "0.788",
        "0.663",
        "Primary model.",
    ],
    [
        "CRCNS+EMS MEM, dropout 0.4",
        "fine-tune",
        "5",
        "0.804 +/- 0.043",
        "0.719 +/- 0.049",
        "0.713",
        "0.725",
        "Higher specificity, lower sensitivity.",
    ],
    [
        "CRCNS+EMS MEM, dropout 0.3",
        "frozen",
        "5",
        "0.719 +/- 0.072",
        "0.613 +/- 0.098",
        "0.813",
        "0.413",
        "Frozen probing is weak.",
    ],
    [
        "CRCNS+EMS MEM, dropout 0.4",
        "frozen",
        "5",
        "0.709 +/- 0.067",
        "0.631 +/- 0.060",
        "0.763",
        "0.500",
        "Frozen probing remains secondary.",
    ],
]


def main() -> None:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_document(doc)

    add_title(doc)
    add_section(
        doc,
        "Abstract",
        [
            (
                "This working draft summarizes the current EyeNet content-agnostic eye-movement "
                "screening pipeline. The model uses fixation-event sequences rather than image, video, "
                "AOI, or task-content features. EMS is the current schizophrenia versus healthy-control "
                "downstream benchmark. Public datasets are first converted into a shared fixation-event "
                "schema and then used for masked event modeling (MEM) pretraining."
            ),
            (
                "Across five EMS stratified splits, EMS-only MEM followed by supervised fine-tuning "
                "improves over from-scratch encoder training. After screening HBN, GazeBase, OneStop, "
                "and CRCNS eye-1, the current primary aligned model is a BiGRU64 encoder pretrained on "
                "EMS+CRCNS eye-1 with span masking probability 0.45 and dropout 0.3, then fine-tuned "
                "on EMS. The evidence remains preliminary because EMS has only 160 labeled subjects."
            ),
        ],
    )

    add_section(
        doc,
        "1. Modeling Boundary",
        [
            (
                "The project is designed for low-burden risk screening and representation learning, "
                "not clinical diagnosis. It intentionally excludes stimulus semantics and dataset-specific "
                "content fields from the universal encoder. Disease labels from different disorders are "
                "not merged into one binary target."
            )
        ],
    )
    add_bullets(
        doc,
        [
            "Universal input unit: fixation event.",
            "Universal feature schema: encoder_no_position_core, 13 features.",
            "Primary downstream target: EMS SZ/HC classification.",
            "Public datasets are used for self-supervised pretraining unless a dataset-specific auxiliary head is explicitly defined.",
        ],
    )

    add_section(doc, "2. Dataset Status", [])
    add_table(
        doc,
        ["Dataset", "Status", "Subjects", "Fixation events", "Role"],
        DATASET_ROWS,
        [1150, 1100, 1050, 1300, 4640],
    )

    add_section(
        doc,
        "3. Current Encoder",
        [
            (
                "The primary model is deliberately small: a 13-dimensional event vector is projected to "
                "64 dimensions, passed through a one-layer bidirectional GRU, pooled by masked attention, "
                "and fine-tuned with a supervised binary head. MEM uses span masking and a learnable mask "
                "token rather than replacing masked values with zero."
            )
        ],
    )
    add_table(
        doc,
        ["Component", "Setting"],
        [
            ["Input", "[batch, sequence_length, 13]"],
            ["Projection", "Linear(13 -> 64) + LayerNorm + ReLU + Dropout"],
            ["Temporal encoder", "1-layer BiGRU, hidden_dim=64, event embedding=128"],
            ["Pooling", "Masked attention pooling"],
            ["Optimizer", "AdamW, learning_rate=1e-3, weight_decay=1e-4"],
            ["Training protocol", "batch_size=8, max_seq_len=1500, gradient_clip_norm=5.0"],
            ["MEM masking", "span masking, probability 0.45 in current primary model, span length 2-8"],
        ],
        [2300, 7060],
    )

    add_section(
        doc,
        "4. EMS Multi-Seed Baseline",
        [
            (
                "The first controlled comparison used five EMS subject-level 60/20/20 splits. "
                "The threshold reported below is selected on the validation split by balanced accuracy."
            )
        ],
    )
    add_table(
        doc,
        ["Experiment", "AUC", "Balanced accuracy", "Sensitivity", "Specificity", "Interpretation"],
        EMS_BASELINE_ROWS,
        [1450, 1300, 1550, 1000, 1000, 3060],
    )

    add_section(
        doc,
        "5. Public Dataset Screening",
        [
            (
                "Public data did not behave as a monotonic scaling variable. CRCNS eye-1 produced the "
                "most promising single-source public transfer result, while GazeBase provided a more "
                "conservative high-specificity operating point. HBN and OneStop are technically integrated "
                "but are not current priority pretraining sources."
            )
        ],
    )
    add_table(
        doc,
        ["Pretraining source", "AUC", "Balanced accuracy", "Sensitivity", "Specificity", "Interpretation"],
        PUBLIC_SCREENING_ROWS,
        [2100, 850, 1350, 1000, 1000, 3060],
    )

    add_section(
        doc,
        "6. Strict Aligned Model Selection",
        [
            (
                "For final model selection, mixed self-supervised pretraining is aligned to the EMS "
                "downstream split: EMS test subjects are not allowed into MEM train. Non-EMS datasets are "
                "split independently by subject. This avoids optimistic leakage from target-dataset test "
                "subjects during self-supervised pretraining."
            )
        ],
    )
    add_table(
        doc,
        ["Experiment", "Mode", "Seeds", "AUC", "Balanced accuracy", "Sensitivity", "Specificity", "Decision"],
        ALIGNED_MODEL_SELECTION_ROWS,
        [1900, 850, 650, 1100, 1350, 900, 900, 1710],
    )

    add_section(
        doc,
        "7. Current Decision",
        [
            (
                "The current primary model is bigru64_ems_crcns_mask045_aligned: BiGRU64, span masking "
                "probability 0.45, dropout 0.3, batch size 8, max sequence length 1500, EMS-anchor "
                "aligned split, and supervised fine-tuning. Dropout 0.4 is retained as a secondary "
                "high-specificity candidate but is not the primary model."
            )
        ],
    )
    add_bullets(
        doc,
        [
            "Do not use frozen probing as the main downstream method.",
            "Do not include single-seed mask0.30 or batch16 runs in the final comparison.",
            "Report multi-seed mean and standard deviation rather than the best single split.",
            "Treat all current EMS results as preliminary because the labeled cohort is small.",
        ],
    )

    add_section(
        doc,
        "8. Next Work",
        [
            (
                "The next stage should pause broad architecture search and focus on documentation, "
                "protocol cleanup, and one targeted confirmation experiment only if needed."
            )
        ],
    )
    add_numbered(
        doc,
        [
            "Finalize the combined result table comparing from-scratch, EMS-only MEM, and aligned CRCNS+EMS MEM.",
            "Document the aligned split protocol as the main no-leakage model-selection protocol.",
            "Optionally test max_seq_len=3000 for the selected dropout 0.3 model while keeping batch size fixed at 8.",
            "Move to additional dataset work only after the current result is cleanly documented.",
        ],
    )

    doc.save(OUTPUT)
    print(OUTPUT)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing = 1.12
    normal.paragraph_format.space_after = Pt(5)

    title = doc.styles["Title"]
    title.font.name = "Calibri"
    title.font.size = Pt(20)
    title.font.bold = True
    title.font.color.rgb = RGBColor(0x0B, 0x25, 0x45)

    for style_name, size, color in [
        ("Heading 1", 15, RGBColor(0x2E, 0x74, 0xB5)),
        ("Heading 2", 12.5, RGBColor(0x2E, 0x74, 0xB5)),
        ("Heading 3", 11.5, RGBColor(0x1F, 0x4D, 0x78)),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(9)
        style.paragraph_format.space_after = Pt(4)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("EyeNet working draft | 2026-05-25")
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)


def add_title(doc: Document) -> None:
    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("Content-Agnostic Eye-Movement Representation Learning for Screening")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Working draft: EMS downstream validation and public-data MEM pretraining")
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)


def add_section(doc: Document, heading: str, paragraphs: list[str]) -> None:
    doc.add_heading(heading, level=1)
    for text in paragraphs:
        doc.add_paragraph(text)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(3)
        paragraph.add_run(item)


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        paragraph = doc.add_paragraph(style="List Number")
        paragraph.paragraph_format.space_after = Pt(3)
        paragraph.add_run(item)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_geometry(table, widths)

    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = header
        set_cell_shading(cell, "EAF2F8")
        format_cell(cell, bold=True, centered=True)

    for row in rows:
        cells = table.add_row().cells
        for idx, text in enumerate(row):
            cells[idx].text = text
            format_cell(cells[idx], bold=False, centered=(idx > 0 and len(text) < 18))

    doc.add_paragraph()


def set_table_geometry(table, widths: list[int]) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_grid = tbl.tblGrid
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        tbl_grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[idx])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_margins(cell, top: int = 80, start: int = 100, bottom: int = 80, end: int = 100) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def format_cell(cell, *, bold: bool, centered: bool) -> None:
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_after = Pt(0)
        if centered:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            run.bold = bold
            run.font.size = Pt(8.0)


if __name__ == "__main__":
    main()
